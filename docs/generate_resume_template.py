"""
生成极简黑白风格简历模板（docx格式）
要求：黑白配色、极简风格、内边距小
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_margins(cell, top=0, bottom=0, left=0, right=0):
    """设置单元格内边距"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin_name, margin_value in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        margin = OxmlElement(f'w:{margin_name}')
        margin.set(qn('w:w'), str(int(margin_value * 20)))  # 转换为twips (1/20 point)
        margin.set(qn('w:type'), 'dxa')
        tcMar.append(margin)
    tcPr.append(tcMar)

def create_resume_template():
    """创建极简黑白风格简历模板"""
    doc = Document()
    
    # 设置页面边距（小内边距）
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # ==================== 个人信息区域 ====================
    # 姓名（大标题）
    name_para = doc.add_paragraph()
    name_run = name_para.add_run('姓名')
    name_run.font.size = Pt(24)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(0, 0, 0)  # 纯黑
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 联系信息（小字，一行）
    contact_para = doc.add_paragraph()
    contact_text = '电话：13800138000  |  邮箱：your.email@example.com  |  地址：XX省XX市'
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(9)
    contact_run.font.color.rgb = RGBColor(0, 0, 0)
    contact_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 分隔线
    doc.add_paragraph('─' * 80).runs[0].font.size = Pt(8)
    doc.add_paragraph()  # 空行
    
    # ==================== 教育背景 ====================
    edu_heading = doc.add_paragraph()
    edu_run = edu_heading.add_run('教育背景')
    edu_run.font.size = Pt(12)
    edu_run.font.bold = True
    edu_run.font.color.rgb = RGBColor(0, 0, 0)
    edu_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 教育经历表格（无边框，小内边距）
    edu_table = doc.add_table(rows=1, cols=3)
    edu_table.style = 'Light Grid Accent 1'
    # 移除表格边框
    for row in edu_table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=50, bottom=50, left=50, right=50)  # 小内边距
            # 移除边框
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'bottom', 'left', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    # 示例数据
    edu_row = edu_table.rows[0]
    edu_row.cells[0].text = '2018.09 - 2022.06'
    edu_row.cells[1].text = 'XX大学'
    edu_row.cells[2].text = '计算机科学与技术 | 本科'
    
    # 设置表格字体
    for row in edu_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].font.size = Pt(10)
                para.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()  # 空行
    
    # ==================== 工作经历 ====================
    work_heading = doc.add_paragraph()
    work_run = work_heading.add_run('工作经历')
    work_run.font.size = Pt(12)
    work_run.font.bold = True
    work_run.font.color.rgb = RGBColor(0, 0, 0)
    work_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # 空行
    
    # ==================== 项目经历 ====================
    project_heading = doc.add_paragraph()
    proj_run = project_heading.add_run('项目经历')
    proj_run.font.size = Pt(12)
    proj_run.font.bold = True
    proj_run.font.color.rgb = RGBColor(0, 0, 0)
    project_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    doc.add_paragraph()  # 空行
    
    # 项目1示例
    proj1_title = doc.add_paragraph()
    proj1_title_run = proj1_title.add_run('1. CheckByAi｜基于双Agent架构的智能物流数据校验系统')
    proj1_title_run.font.size = Pt(11)
    proj1_title_run.font.bold = True
    proj1_title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    proj1_info = doc.add_paragraph()
    proj1_info_run = proj1_info.add_run('项目角色：全栈开发（Agent架构核心负责人）  |  项目周期：2025.1 – 2025.2')
    proj1_info_run.font.size = Pt(9)
    proj1_info_run.font.color.rgb = RGBColor(0, 0, 0)
    
    proj1_tech = doc.add_paragraph()
    proj1_tech_run = proj1_tech.add_run('技术栈：Spring Boot 3.2、Vue 3、LangGraph、LangChain、Coze工作流、PaddleOCR、DeepSeek LLM、MyBatis-Plus、Redis、JWT、MySQL、Nginx')
    proj1_tech_run.font.size = Pt(9)
    proj1_tech_run.font.color.rgb = RGBColor(0, 0, 0)
    
    proj1_desc = doc.add_paragraph()
    proj1_desc_run = proj1_desc.add_run('项目描述：')
    proj1_desc_run.font.size = Pt(10)
    proj1_desc_run.font.bold = True
    proj1_desc_run.font.color.rgb = RGBColor(0, 0, 0)
    
    proj1_desc_text = doc.add_paragraph()
    proj1_desc_text_run = proj1_desc_text.add_run('设计并实现双Agent架构的智能物流数据校验系统，支持Coze工作流（低代码平台）与LangGraph多Agent协作两种处理模式，用户可通过前端开关灵活切换。系统采用前后端分离架构，支持批量数据处理、异步回调机制、任务状态实时追踪，处理准确率>95%。')
    proj1_desc_text_run.font.size = Pt(10)
    proj1_desc_text_run.font.color.rgb = RGBColor(0, 0, 0)
    
    proj1_highlights = doc.add_paragraph()
    proj1_highlights_run = proj1_highlights.add_run('核心工作与亮点：')
    proj1_highlights_run.font.size = Pt(10)
    proj1_highlights_run.font.bold = True
    proj1_highlights_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 亮点列表
    highlights = [
        '设计并实现双Agent架构，集成Coze工作流与LangGraph多Agent协作两种处理模式，前端提供可视化开关实现运行时动态切换。',
        '实现LangGraph多Agent协作架构，构建Coordinator、Excel、OCR、Think四个专业化Agent，通过StateGraph实现状态机工作流与智能路由决策。',
        '优化OCR批量处理流程，采用独立UUID目录策略解决多行并发处理时的文件冲突问题，确保并发安全与资源隔离。',
        '设计专业化LLM Prompt工程，实现4层校验规则，基于DeepSeek LLM进行重量一致性智能校验，生成结构化分析报告。'
    ]
    
    for i, highlight in enumerate(highlights, 1):
        highlight_para = doc.add_paragraph()
        highlight_run = highlight_para.add_run(f'{i}. {highlight}')
        highlight_run.font.size = Pt(10)
        highlight_run.font.color.rgb = RGBColor(0, 0, 0)
        # 设置首行缩进
        highlight_para.paragraph_format.left_indent = Inches(0.2)
    
    doc.add_paragraph()  # 空行
    
    # ==================== 技能专长 ====================
    skills_heading = doc.add_paragraph()
    skills_run = skills_heading.add_run('技能专长')
    skills_run.font.size = Pt(12)
    skills_run.font.bold = True
    skills_run.font.color.rgb = RGBColor(0, 0, 0)
    skills_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    skills_table = doc.add_table(rows=1, cols=2)
    # 移除表格边框
    for row in skills_table.rows:
        for cell in row.cells:
            set_cell_margins(cell, top=30, bottom=30, left=30, right=30)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'bottom', 'left', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'nil')
                tcBorders.append(border)
            tcPr.append(tcBorders)
    
    skills_row = skills_table.rows[0]
    skills_row.cells[0].text = '编程语言：Java、Python、JavaScript'
    skills_row.cells[1].text = '框架技术：Spring Boot、Vue 3、LangGraph、LangChain'
    
    for row in skills_table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                para.runs[0].font.size = Pt(10)
                para.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()  # 空行
    
    # ==================== 其他信息 ====================
    other_heading = doc.add_paragraph()
    other_run = other_heading.add_run('其他信息')
    other_run.font.size = Pt(12)
    other_run.font.bold = True
    other_run.font.color.rgb = RGBColor(0, 0, 0)
    other_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    other_para = doc.add_paragraph()
    other_text = 'GitHub：https://github.com/yourusername  |  个人博客：https://yourblog.com'
    other_run = other_para.add_run(other_text)
    other_run.font.size = Pt(9)
    other_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 保存文档
    output_path = 'resume_template.docx'
    doc.save(output_path)
    print(f'简历模板已生成：{output_path}')
    print('提示：请使用Word打开文件，根据实际情况修改内容')

if __name__ == '__main__':
    try:
        from docx import Document
        create_resume_template()
    except ImportError:
        print('错误：需要安装 python-docx 库')
        print('安装命令：pip install python-docx')
        print('\n或者运行以下命令安装依赖：')
        print('pip install python-docx')

